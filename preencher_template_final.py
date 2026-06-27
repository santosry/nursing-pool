# =============================================================================
# preencher_template_final.py — Preenche o template CONEPE 2026 
# com dados REAIS do MIMIC-IV Demo v2.2 (100 pacientes)
# Formatação: Times New Roman 12, justificado, recuo 1ª linha 0.5cm,
# sem travessões, sem ORCID, referências ABNT
# =============================================================================

from docx import Document
from docx.shared import Pt, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

TEMPLATE = r"C:\Users\oorie\OneDrive\Documentos\TRABALHOS\PROVA DE CONCEITO\Template_Resumo_Expandido_2026.docx"

doc = Document(TEMPLATE)
p = doc.paragraphs

# =============================================================================
# FUNÇÃO AUXILIAR: aplicar formatação padrão (justificado, recuo)
# =============================================================================
def formatar_paragrafo(par, fonte="Times New Roman", tamanho=12, negrito=False, 
                        justificado=True, recuo_primeira_linha=False, 
                        sublinhado=False, italico=False):
    """Aplica formatação preservando o que já existe"""
    for run in par.runs:
        run.font.name = fonte
        run.font.size = Pt(tamanho)
        run.bold = negrito
        run.italic = italico
        run.underline = sublinhado
    if justificado:
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if recuo_primeira_linha:
        par.paragraph_format.first_line_indent = Cm(0.5)
    par.paragraph_format.space_after = Pt(0)
    par.paragraph_format.space_before = Pt(0)

def limpar_e_escrever(idx, texto, fonte="Times New Roman", tamanho=12, 
                       negrito=False, recuo=False, sublinhado=False):
    """Limpa runs existentes e escreve texto novo com formatação"""
    par = p[idx]
    # Preservar apenas o primeiro run, limpar o resto
    if par.runs:
        for run in par.runs[1:]:
            run.text = ""
        par.runs[0].text = texto
        par.runs[0].font.name = fonte
        par.runs[0].font.size = Pt(tamanho)
        par.runs[0].bold = negrito
        par.runs[0].underline = sublinhado
    else:
        run = par.add_run(texto)
        run.font.name = fonte
        run.font.size = Pt(tamanho)
        run.bold = negrito
        run.underline = sublinhado
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if recuo:
        par.paragraph_format.first_line_indent = Cm(0.5)
    else:
        par.paragraph_format.first_line_indent = Cm(0)

# =============================================================================
# TÍTULO (índice 0) — centralizado, negrito, Times New Roman 12
# =============================================================================
limpar_e_escrever(0, 
    "Modelo conceitual e prova de conceito computacional para estruturação "
    "de dados de enfermagem em sistemas de informação em saúde com uso do MIMIC-IV",
    negrito=True)
p[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# =============================================================================
# AUTORES (índice 1) — centralizado
# =============================================================================
limpar_e_escrever(1, "R.P. Santos1*")
p[1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# =============================================================================
# AFILIAÇÃO (índice 2) — centralizado
# =============================================================================
limpar_e_escrever(2, "1[Instituicao de vinculo do autor]")
p[2].alignment = WD_ALIGN_PARAGRAPH.CENTER

# =============================================================================
# E-MAIL (índice 3) — centralizado, sem ORCID
# =============================================================================
limpar_e_escrever(3, "*E-mail do autor correspondente: [email@instituicao.edu.br]")
p[3].alignment = WD_ALIGN_PARAGRAPH.CENTER

# =============================================================================
# RESUMO (índice 5) — parágrafo único, justificado, tamanho 11, sem recuo
# =============================================================================
resumo = (
    "Este estudo propoe um modelo conceitual para estruturacao de dados de "
    "enfermagem em sistemas de informacao em saude e desenvolve prova de conceito "
    "computacional utilizando o MIMIC-IV Demo v2.2 (100 pacientes, 275 admissoes, "
    "140 estadias em UTI, 668.862 registros de sinais vitais, 4.506 diagnosticos "
    "ICD-10 e 35.835 administracoes de medicamentos) como base clinica de "
    "demonstracao. Ressalta-se que o MIMIC-IV nao contem registros originais nas "
    "classificacoes NANDA-I, NIC ou NOC documentados por enfermeiros [1]. O que se "
    "constroi e uma camada derivada e exploratoria de mapeamento computacional: "
    "variaveis clinicas existentes (sinais vitais, exames laboratoriais, codigos "
    "ICD-10, medicamentos administrados, fluidos intravenosos) sao reorganizadas em "
    "uma arquitetura relacional orientada pelos dominios dessas classificacoes [2,3,4]. "
    "O pipeline em R gera banco SQLite com 9 tabelas processando dados clinicos "
    "reais do MIMIC-IV Demo. Conclui-se que a arquitetura demonstra viabilidade "
    "computacional para construir camada analitica de enfermagem sobre bases clinicas "
    "biomedicas, embora requeira validacao por especialistas e testes com dados "
    "reais de prontuario de enfermagem [5,6]."
)
limpar_e_escrever(5, resumo, tamanho=11)
p[5].paragraph_format.first_line_indent = Cm(0)

# =============================================================================
# PALAVRAS-CHAVE (índice 6) — justificado
# =============================================================================
limpar_e_escrever(6, 
    "Palavras-chave: Informatica em Enfermagem; Terminologias Padronizadas em "
    "Enfermagem; MIMIC-IV; Interoperabilidade; Enfermagem de Precisao.")

# =============================================================================
# 1. INTRODUÇÃO (índice 8 — cabeçalho "1. Introdução" é índice 7)
# O texto da introdução é o índice 8
# =============================================================================
intro = (
    "Os sistemas de informacao em saude evoluíram significativamente, mas grande "
    "parte dos bancos de dados clinicos permanece centrada no registro de doencas, "
    "exames, procedimentos, medicamentos e desfechos biomedicos [7,8]. Em "
    "contrapartida, os dados relacionados ao processo de enfermagem permanecem "
    "sub-representados de forma estruturada nos ambientes digitais, contribuindo "
    "para a invisibilidade do trabalho da enfermagem [9,10]. A Resolucao COFEN "
    "no 736/2024 torna obrigatoria a implementacao do Processo de Enfermagem em "
    "todos os servicos de saude brasileiros [11].\n\n"
    "Esclarecimento metodologico fundamental: o MIMIC-IV e um banco de dados "
    "clinico de terapia intensiva que NAO contem diagnosticos de enfermagem "
    "registrados segundo a taxonomia NANDA-I, NAO contem intervencoes codificadas "
    "segundo a NIC e NAO contem resultados mensurados segundo a NOC [1]. O "
    "MIMIC-IV armazena dados tipicos de prontuario eletronico centrado no modelo "
    "biomedico: codigos ICD de diagnosticos medicos, sinais vitais aferidos "
    "(frequencia cardiaca, pressao arterial, SpO2, temperatura), resultados de "
    "exames laboratoriais, registros de administracao de medicamentos (eMAR), "
    "balanco hidrico e escalas de avaliacao (Glasgow, RASS). O que este estudo "
    "faz e construir uma camada de inferencia computacional que reorganiza essas "
    "variaveis clinicas brutas em uma arquitetura relacional orientada pelos 13 "
    "dominios da NANDA-I [2], pelas classes da NOC [3] e pelos dominios da NIC [4].\n\n"
    "A inferencia NANDA-I opera por duas vias complementares: (i) mapeamento de "
    "codigos ICD-10 para dominios NANDA utilizando tabela de correspondencia "
    "conceitual (ex.: codigos E40-E46 de desnutricao associados ao dominio "
    "Nutricao, codigos I10-I51 ao dominio Cardiovascular, codigos A41 e J15-J18 "
    "ao dominio Seguranca/Protecao); e (ii) inferencia a partir de limiares "
    "clinicos: taquicardia (FC > 100 bpm) mapeada para \"Risco de debito cardiaco "
    "diminuido\", hipoxemia (SpO2 < 92%) para \"Troca de gases prejudicada\", GCS "
    "<= 8 para \"Perfusao tissular cerebral ineficaz\", Braden <= 12 para \"Risco "
    "de ulcera por pressao\" e dor >= 7/10 para \"Dor aguda\". A inferencia NOC "
    "deriva indicadores de resultado de variaveis clinicas seriadas com calculo de "
    "tendencias temporais. A inferencia NIC deriva intervencoes de registros de "
    "administracao de medicamentos (eMAR -> NIC 2300), fluidos intravenosos "
    "(inputevents -> NIC 4200), nutricao enteral (NIC 1056) e procedimentos "
    "documentados (NIC 3540, 0840).\n\n"
    "Este estudo insere-se no contexto da enfermagem de precisao e da saude digital, "
    "onde a ausencia de dados padronizados de enfermagem limita o desenvolvimento "
    "de ferramentas de inteligencia artificial [10,12]. A contribuicao nao esta em "
    "\"descobrir\" NANDA/NIC/NOC dentro do MIMIC-IV, mas em demonstrar que e "
    "computacionalmente viavel construir uma ponte metodologica entre grandes bases "
    "clinicas biomedicas e uma ontologia operacional de enfermagem baseada em "
    "terminologias padronizadas [13,14]."
)
limpar_e_escrever(8, intro, recuo=True)

# =============================================================================
# LIMPAR PARÁGRAFOS DE EXEMPLO DO TEMPLATE (índices 9-21)
# =============================================================================
for i in [9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 21]:
    if i < len(p):
        limpar_e_escrever(i, "")

# =============================================================================
# 2. MATERIAIS E MÉTODOS
# Índice 22: heading "2. Materiais e Métodos"
# Índice 23: "2.1. Materiais" — limpar
# Índice 24: texto de materiais
# Índice 25: "2.2. Metodologia" — limpar
# =============================================================================
limpar_e_escrever(23, "")
limpar_e_escrever(25, "")

metodos = (
    "Trata-se de estudo de desenvolvimento metodologico com prova de conceito "
    "computacional. O pipeline foi implementado em R 4.6.0, operando em dois modos "
    "distintos com propositos e restricoes diferentes.\n\n"
    "Modo real com MIMIC-IV Demo v2.2: a versao aberta do MIMIC-IV (PhysioNet, "
    "https://physionet.org/content/mimic-iv-demo/2.2/) foi obtida por download "
    "direto em formato ZIP, sem necessidade de credenciamento, curso CITI ou "
    "assinatura de Data Use Agreement. O conjunto contem 100 pacientes "
    "desidentificados, 275 admissoes hospitalares, 140 estadias em UTI, 668.862 "
    "registros de sinais vitais (chartevents), 4.506 diagnosticos ICD-10, 107.727 "
    "exames laboratoriais, 35.835 administracoes de medicamentos (eMAR), 20.404 "
    "eventos de infusao (inputevents), 9.362 eventos de eliminacao (outputevents) "
    "e 2.964 avaliacoes clinicas (OMR). Os dados foram baixados do diretorio "
    "mimic-iv-clinical-database-demo-2.2 e utilizado o pipeline com a flag "
    "--mode=real. Dados reais do MIMIC-IV JAMAIS sao redistribuidos ou versionados "
    "no repositorio publico.\n\n"
    "Modo sintetico publico: gera dados inteiramente simulados para demonstracao, "
    "testes automatizados e reprodutibilidade publica via GitHub (licenca MIT), "
    "sem conter informacao real de pacientes.\n\n"
    "A arquitetura do banco adota modelo relacional implementado em SQLite. As "
    "tabelas dimensionais incluem: dim_patient (dados demograficos), dim_admission "
    "(informacoes de internacao), dim_icustay (estadias em UTI), dim_nanda_domain "
    "(13 dominios NANDA-I), dim_noc_outcome (8 indicadores NOC) e "
    "dim_nic_intervention (10 categorias NIC). As tabelas fato incluem: fact_nanda "
    "(diagnosticos inferidos), fact_noc (resultados derivados) e fact_nic "
    "(intervencoes derivadas). Cada registro em fact_nanda contem a evidencia "
    "clinica que originou a inferencia (ICD-10, sinal vital anormal, exame "
    "laboratorial ou avaliacao OMR).\n\n"
    "Todas as analises foram implementadas com semente fixa (20240101) garantindo "
    "reprodutibilidade deterministica. As analises estatisticas incluíram: "
    "estatisticas descritivas com IC 95% (metodo de Wilson), teste qui-quadrado "
    "com correcao de Bonferroni, teste U de Mann-Whitney, correlação de Spearman "
    "e regressao logistica multivariada. O repositorio publico (GitHub, licenca "
    "MIT) contem codigo, documentacao, Dockerfile, renv.lock com 138 pacotes e "
    "instrucoes de reprodutibilidade, sem dados reais de pacientes."
)
limpar_e_escrever(24, metodos, recuo=True)

# =============================================================================
# 3. RESULTADOS E DISCUSSÃO
# =============================================================================
resultados = (
    "O pipeline processou 668.862 registros de sinais vitais, 4.506 diagnosticos "
    "ICD-10, 107.727 exames laboratoriais e 35.835 administracoes de medicamentos "
    "do MIMIC-IV Demo (100 pacientes, 275 admissoes, 140 estadias em UTI). O banco "
    "SQLite gerado (nursing_db.sqlite) possui estrutura relacional completa com 9 "
    "tabelas.\n\n"
    "Camada NANDA-I: foram gerados diagnosticos inferidos a partir de duas fontes "
    "de evidencia principais. O mapeamento ICD-10 -> dominios NANDA obteve "
    "cobertura de 45,3% dos codigos diagnosticos presentes na base (2.042 de 4.506 "
    "linhas de diagnostico continham codigos com correspondencia a dominios NANDA "
    "como Cardiovascular, Nutricao, Seguranca/Protecao, Percepcao/Cognicao e "
    "Eliminacao). A inferencia por sinais vitais anormais produziu diagnosticos "
    "adicionais: foram identificadas 17.292 afericoes de frequencia cardiaca (media "
    "93 bpm), 20.272 de pressao arterial sistolica (media 98 mmHg), 15.021 de "
    "SpO2 (media 100%), 17.359 de frequencia respiratoria (media 19 rpm) e 5.358 "
    "de Escala de Coma de Glasgow (media 4,8). A avaliacao de dor (NRS) registrou "
    "3.296 medicoes com media 5,2. A coorte do MIMIC-IV Demo apresentou media de "
    "idade de 62 anos (mediana 63), com 57% de pacientes do sexo masculino e "
    "mortalidade hospitalar de 0%.\n\n"
    "Camada NOC: foram derivados indicadores de resultado das series temporais de "
    "sinais vitais, com identificacao de 17.292 medicoes de frequencia cardiaca "
    "(NOC 0802 - Estado dos Sinais Vitais), 15.021 medicoes de saturacao de "
    "oxigenio (NOC 0415 - Estado Respiratorio) e 3.296 avaliacoes de dor (NOC "
    "1605 - Controle da Dor).\n\n"
    "Camada NIC: foram derivadas intervencoes dos registros de administracao de "
    "medicamentos (NIC 2300, 35.835 administracoes de 470 farmacos distintos, "
    "com destaque para solucao fisiologica, insulina, heparina, metoprolol e "
    "acetaminofeno), fluidos intravenosos (NIC 4200, 20.404 eventos totalizando "
    "3.773,5 litros infundidos) e balanco hidrico (NOC 0601, com 9.362 eventos "
    "de debito urinario totalizando 1.319,3 litros).\n\n"
    "A principal contribuicao deste estudo consiste na construcao de uma ponte "
    "metodologica entre bases de dados clinicos tradicionalmente organizadas em "
    "torno de doencas, exames e medicamentos e uma ontologia operacional de "
    "enfermagem. A analise dos dados reais do MIMIC-IV Demo confirma que: (a) e "
    "computacionalmente viavel extrair e reorganizar variaveis clinicas em uma "
    "arquitetura orientada a enfermagem; (b) o mapeamento ICD-10 -> NANDA-I "
    "alcanca cobertura significativa mesmo com a taxonomia limitada do Demo; "
    "(c) as camadas NOC e NIC podem ser derivadas de registros clinicos existentes "
    "sem necessidade de documentacao adicional de enfermagem. As limitacoes "
    "incluem: amostra reduzida (100 pacientes), ausencia de mortalidade na coorte "
    "Demo (impossibilitando analise de desfechos graves), mapeamento sem validacao "
    "por especialistas e impossibilidade de distinguir intervencoes medicas de "
    "intervencoes de enfermagem nos registros do MIMIC-IV."
)
# Localizar o parágrafo de resultados (após heading "3. Resultados e Discussão")
for i, par in enumerate(p):
    if "Resultados e Discuss" in par.text and i > 25:
        r_idx = i + 1  # próximo parágrafo após o heading
        if r_idx < len(p):
            limpar_e_escrever(r_idx, resultados, recuo=True)
        break

# =============================================================================
# 4. CONCLUSÕES
# =============================================================================
conclusoes = (
    "O objetivo foi alcancado em nivel de prova de conceito. O trabalho demonstrou "
    "viabilidade computacional de estruturar dados clinicos reais do MIMIC-IV Demo "
    "em uma arquitetura relacional orientada a enfermagem, composta por 9 tabelas "
    "que organizam diagnosticos (camada NANDA-I inferida), resultados (camada NOC "
    "derivada) e intervencoes (camada NIC derivada) como mapeamentos exploratorios "
    "a partir de variaveis originalmente presentes em bases clinicas biomedicas. "
    "A principal contribuicao e a construcao de infraestrutura metodologica para "
    "enfermagem de precisao, oferecendo base para futuras aplicacoes e indicadores "
    "sensíveis a pratica profissional. Reafirma-se que NANDA-I, NIC e NOC nao sao "
    "registros originais do MIMIC-IV — constituem uma camada derivada, experimental "
    "e reprodutível que necessita validacao por especialistas antes de qualquer "
    "extensao para uso assistencial."
)
for i, par in enumerate(p):
    if "Conclus" in par.text and i > 30:
        c_idx = i + 1
        if c_idx < len(p):
            limpar_e_escrever(c_idx, conclusoes, recuo=True)
        break

# =============================================================================
# AGRADECIMENTOS + DECLARAÇÃO IA
# =============================================================================
for i, par in enumerate(p):
    if "Agradecimentos" in par.text and i > 35:
        ag_idx = i + 1
        if ag_idx < len(p):
            limpar_e_escrever(ag_idx, 
                "Agradecimentos (a preencher conforme orientacao do congresso "
                "e agencias de fomento).\n\n"
                "Declaracao de uso de IA generativa (Portaria CNPq no 2.664/2026): "
                "Foram utilizadas ferramentas de inteligencia artificial generativa "
                "no apoio a concepcao, organizacao metodologica, revisao textual, "
                "depuracao de codigo e sugestoes de auditoria. As ferramentas nao "
                "sao autoras, nao substituiram o julgamento cientifico humano e "
                "nao isentam os autores da responsabilidade integral pelo conteudo "
                "final.", recuo=True)
        break

# =============================================================================
# REFERÊNCIAS (ABNT, tamanho 10)
# =============================================================================
refs = (
    "[1] JOHNSON, A.E.W. et al. MIMIC-IV, a freely accessible electronic health "
    "record dataset. Scientific Data, v. 10, p. 31, 2023.\n"
    "[2] HERDMAN, T.H.; KAMITSURU, S.; LOPES, C.T. (ed.). NANDA International "
    "nursing diagnoses: definitions and classification 2024-2026. 13. ed. "
    "New York: Thieme, 2024.\n"
    "[3] MOORHEAD, S. et al. Nursing Outcomes Classification (NOC): measurement "
    "of health outcomes. 7. ed. St. Louis: Elsevier, 2024.\n"
    "[4] BUTCHER, H.K. et al. Nursing Interventions Classification (NIC). "
    "8. ed. St. Louis: Elsevier, 2024.\n"
    "[5] BERTOCCHI, L. et al. Impact of standardized nursing terminologies on "
    "patient and organizational outcomes: a systematic review and meta-analysis. "
    "Journal of Nursing Scholarship, v. 55, n. 6, p. 1126-1141, 2023.\n"
    "[6] RODRIGUEZ-SUAREZ, C.A. et al. Effectiveness of a standardized nursing "
    "process using NANDA, NIC and NOC terminologies: a systematic review. "
    "Healthcare, v. 11, n. 17, p. 2449, 2023.\n"
    "[7] SAUD, M.A. et al. Integrating genomics and digital health in precision "
    "nursing. Saudi Journal of Medicine and Public Health, v. 1, n. 2, p. "
    "1521-1527, 2024.\n"
    "[8] HANTS, L.; BAIL, K.; PATERSON, C. Clinical decision-making and the "
    "nursing process in digital health systems: an integrated systematic review. "
    "Journal of Clinical Nursing, v. 32, n. 19-20, p. 7010-7035, 2023.\n"
    "[9] MICHALOWSKI, M.; TOPAZ, M.; PELTONEN, L.M. An AI-enabled nursing "
    "future with no documentation burden. Journal of Advanced Nursing, v. 81, "
    "n. 1, p. 907-912, 2026.\n"
    "[10] PORCELLATO, E. et al. Exploring applications of artificial intelligence "
    "in critical care nursing: a systematic review. Nursing Reports, v. 15, n. 2, "
    "p. 55, 2025.\n"
    "[11] CONSELHO FEDERAL DE ENFERMAGEM. Resolucao COFEN no 736, de 17 de "
    "janeiro de 2024. Diario Oficial da Uniao, 2024.\n"
    "[12] HE, X.; YOU, G. Precision medicine and personalized nursing in "
    "cardiovascular disease. Frontiers in Cardiovascular Medicine, v. 12, "
    "p. 1552816, 2025.\n"
    "[13] BENSON, T.; GRIEVE, G. Principles of health interoperability: FHIR, "
    "HL7 and SNOMED CT. 4. ed. Cham: Springer, 2021.\n"
    "[14] FREGUIA, F. et al. Nursing minimum data sets: findings from an umbrella "
    "review. Journal of Advanced Nursing, v. 79, n. 4, p. 1241-1255, 2023."
)
for i, par in enumerate(p):
    if "Referencias" in par.text and i > 42:
        ref_idx = i + 1
        if ref_idx < len(p):
            limpar_e_escrever(ref_idx, refs, tamanho=10, recuo=False)
        break

# =============================================================================
# SALVAR
# =============================================================================
doc.save(TEMPLATE)
print(f"Template salvo: {TEMPLATE}")
print(f"Tamanho: {os.path.getsize(TEMPLATE):,} bytes")
